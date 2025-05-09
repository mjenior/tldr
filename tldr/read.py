
import os

from docx import Document
from PyPDF2 import PdfReader, PdfMerger
from bs4 import BeautifulSoup


def fetch_content(search_dir):
	"""
	Find files and read in thier contents. 
	Returns a dictionary with file extensions as keys.
	"""

	# Collect readable text
	files = _find_readable_files(search_dir)

	# Extract text from found files
	content = {}
	for ext in files.keys():
		content[ext] = _read_file_content(files[ext], ext)

	return content


def _find_readable_files(directory: str = '.') -> dict:
    """
    Recursively scan the given directory for readable text files.
    Includes: .pdf, .docx, and .html files, and generic text files.
    Args:
        directory: The path to the directory to scan. Defaults to the current directory.
    Returns:
        A dictionary with file extensions as keys and lists of file paths as values.
    """
    readable_files_by_type = {'pdf': [], 'docx': [], 'html': [], 'txt': []}

    for root, _, files in os.walk(directory):
        for filename in files:
            filepath = os.path.join(root, filename)
            ext = os.path.splitext(filename)[1].lower()

            try:
                if ext in ['.pdf']:
                    reader = PdfReader(filepath)
                    # Check if any page seems to contain text
                    if any(page.extract_text() for page in reader.pages):
                        readable_files_by_type['pdf'].append(filepath)

                elif ext in ['.docx']:
                    doc = Document(filepath)
                    if any(para.text.strip() for para in doc.paragraphs):
                        readable_files_by_type['docx'].append(filepath)

                elif ext in ['.html', '.htm']:
                    with open(filepath, 'r', encoding='utf-8') as f:
                        f.read(1024)
                    readable_files_by_type['html'].append(filepath)

                else:
                    with open(filepath, 'r', encoding='utf-8') as f:
                        f.read(1024)
                    readable_files_by_type['txt'].append(filepath)

            except Exception as e:
            	#print(f"Error analyzing file '{filepath}': {e}")
                continue

    return {k: v for k, v in readable_files_by_type.items() if v}


 def _read_file_content(filelist, ext):
    """
    Reads the main body text from a given file paths and returns as strings.
    """
    content_list = []
    for ext in filelist:

	    try:
	        if ext == 'pdf':
	            reader = PdfReader(filepath)
	            text = ""
	            for page in reader.pages:
	                text += page.extract_text() + "\n"
	        elif 'doc' in ext:
	            doc = Document(filepath)
	            text = ""
	            for para in doc.paragraphs:
	                text += para.text + "\n"
	        elif 'htm' in ext:
	            with open(filepath, 'r', encoding='utf-8') as f:
	                html_content = f.read()
	            soup = BeautifulSoup(html_content, 'html.parser')
	            text = soup.get_text(separator='\n')
	        else:
	            with open(filepath, 'r', encoding='utf-8') as f:
	                text = f.read()
	    except Exception as e:
	        print(f"Error reading file '{filepath}': {e}")

	    content_list.append(text.strip())

    return content_list


