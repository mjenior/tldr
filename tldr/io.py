import os
import re
import yaml
import string
import unicodedata
from datetime import datetime

from docx import Document
from PyPDF2 import PdfReader
from bs4 import BeautifulSoup

from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.units import inch


class FileHandler:
    """
    Handles all file system operations for the application.

    This class is responsible for finding, reading, and writing files in various
    formats, including text, PDF, DOCX, and YAML. It also manages the creation
    of timestamped output directories for each run, ensuring that intermediate
    and final files are stored in an organized manner.

    Attributes:
        timestamp (str): A timestamp string generated for the current run,
            e.g., '20231027_103000'.
        run_tag (str): A unique identifier for the current run, combining a
            prefix with the timestamp, e.g., 'tldr.20231027_103000'.
        output_directory (str): The path to the directory where output files
            for the current run are stored.
    """

    def _generate_run_tag(self):
        """Generate a timestamp string (e.g., 20231027_103000)"""
        self.timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        self.run_tag = f"tldr.{self.timestamp}"

    def _create_output_path(self):
        """Set up where to write intermediate files"""
        self.output_directory = f"{self.run_tag}.files"
        os.makedirs(self.output_directory, exist_ok=True)

    def fetch_content(
        self,
        label: str = "input",
        user_files: list = None,
        search_dir: str = ".",
        recursive: bool = False,
    ):
        """
        Find files and read in their contents.
        Returns a dictionary of file paths and content strings.
        """
        self.logger.info(f"Reading content of {label} documents...")
        files = self._find_readable_files(
            infiles=user_files, directory=search_dir, recursive=recursive
        )

        # Create content dictionary
        content = {}
        for ext, filelist in files.items():
            for f in filelist:
                f_content = self.read_file_content(f, ext)
                if "readme" in f.lower():
                    f_test = False
                else:
                    f_test = self.is_text_corrupted(f_content)
                if f_test is True:
                    self.logger.error(f"File '{f}' may be corrupted, ignoring.")
                    continue
                elif f_content is not None:
                    content[f] = {"content": f_content}

        return content

    def _find_readable_files(
        self, infiles: list = [], directory: str = ".", recursive: bool = False
    ) -> dict:
        """
        Scan for readable text files.
        Args:
            infiles: Optional list of specific files to process.
            directory: The path to the directory to scan. Defaults to the current directory.
            recursive: Whether to search recursively in subdirectories.
        Returns:
            A dictionary with file extensions as keys and lists of file paths as values.
        """
        readable_files_by_type = {
            "pdf": [],
            "docx": [],
            "html": [],
            "txt": [],
            "md": [],
        }

        # Handle user selected files first
        infiles_found = 0
        for file_path_item in infiles:
            infiles_found += 1
            ext = file_path_item.split(".")[-1].lower().strip()
            self._update_file_dictionary(
                readable_files_by_type, file_path_item, ext
            )
            self.logger.info(f"Found file '{file_path_item}' of type '{ext}'.")
        if infiles_found >= 1:
            return {k: v for k, v in readable_files_by_type.items() if v}

        # Then handle directory search in none are returned
        elif recursive is False:
            for filename in os.listdir(directory):
                if ".tldr." in filename:
                    continue
                filepath = os.path.join(directory, filename)
                if os.path.isfile(filepath):
                    ext = filepath.split(".")[-1].lower().strip()
                    self._update_file_dictionary(
                        readable_files_by_type, filepath, ext
                    )
                    self.logger.info(f"Found file '{filepath}' of type '{ext}'.")
        elif recursive is True:
            for root, _, files_in_dir in os.walk(directory):
                for filename in files_in_dir:
                    if ".tldr." in filename:
                        continue
                    filepath = os.path.join(root, filename)
                    ext = filepath.split(".")[-1].lower().strip()
                    self._update_file_dictionary(
                        readable_files_by_type, filepath, ext
                    )
                    self.logger.info(f"Found file '{filepath}' of type '{ext}'.")

        return {k: v for k, v in readable_files_by_type.items() if v}

    def _update_file_dictionary(self, file_dict, file_path, file_ext):
        """Add file path to dictionary if readable and of a supported type."""
        try:
            if file_ext == "pdf":
                reader = PdfReader(file_path)
                if reader.pages and any(
                    page.extract_text() for page in reader.pages if page.extract_text()
                ):
                    file_dict["pdf"].append(file_path)
            elif file_ext == "docx":
                doc = Document(file_path)
                if any(para.text.strip() for para in doc.paragraphs):
                    file_dict["docx"].append(file_path)
            elif file_ext in ["html", "htm"]:
                with open(file_path, "r", encoding="utf-8") as f:
                    f.read(1024)  # Try reading a small chunk to check readability
                file_dict["html"].append(file_path)
            elif file_ext == "md":
                with open(file_path, "r", encoding="utf-8") as f:
                    f.read(1024)
                file_dict["md"].append(file_path)
            elif file_ext == "txt":
                with open(file_path, "r", encoding="utf-8") as f:
                    f.read(1024)
                file_dict["txt"].append(file_path)
        except Exception as e:  # pylint: disable=broad-except
            # Silently ignore files that can't be processed or raise an error
            pass

    def read_file_content(self, filepath, ext=None):
        """
        Reads the main body text from given file paths and returns as strings.
        """
        current_ext = ext
        if current_ext is None:
            current_ext = filepath.split(".")[-1].lower().strip()

        try:
            if current_ext == "pdf":
                reader = PdfReader(filepath)
                text = ""
                for page in reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
            elif current_ext in ["doc", "docx"]:
                doc = Document(filepath)
                text = "\n".join([para.text for para in doc.paragraphs])
            elif current_ext in ["htm", "html"]:
                with open(filepath, "r", encoding="utf-8") as f:
                    html_content = f.read()
                soup = BeautifulSoup(html_content, "html.parser")
                text = soup.get_text(separator="\n")
            elif current_ext in ["txt", "md"]:
                with open(filepath, "r", encoding="utf-8") as f:
                    text = f.read()

            return text.strip()

        except Exception as e:  # pylint: disable=broad-except
            self.logger.error(f"Error reading file '{filepath}': {e}")
            return None

    def is_text_corrupted(self, text: str, threshold: float = 0.25) -> bool:
        """
        Determines if a text stream is likely corrupted by analyzing its content.

        Args:
            text: The text content to analyze
            threshold: The ratio of suspicious patterns to total characters
                     above which the text is considered corrupted (default: 0.25)

        Returns:
            bool: True if the text is likely corrupted, False otherwise
        """
        if text is None or len(text.strip()) == 0:
            return True

        # Check for binary null bytes which are a strong indicator of binary data
        if '\x00' in text:
            return True

        # Check for common corruption patterns
        corruption_patterns = [
            r'[\x80-\xFF]',  # High ASCII/Unicode garbage
            r'\S{30,}',       # Very long words (reduced from 50)
            r'[0-9]{8,}',     # Long sequences of numbers (reduced from 10)
            r'[^\x00-\x7F]{2,}',  # Multiple consecutive non-ASCII (reduced from 3)
            r'[\uFFFD]',      # Unicode replacement character
            r'[\x00-\x08\x0B\x0C\x0E-\x1F\x7F]',  # Control chars except \t, \n, \r
            r'[\u0000-\u001F\u007F-\u009F]',  # Additional control characters
            r'[\uFFFB-\uFFFD]',  # More replacement characters
            r'\\x[0-9A-Fa-f]{2}',  # Hex-encoded bytes
            r'\\u[0-9A-Fa-f]{4}',  # Unicode escape sequences
            r'[\u0080-\u00FF]',  # Extended ASCII (common in encoding errors)
            r'[\u0000-\u0008\u000B\u000C\u000E-\u001F\u007F-\u009F]',  # More control chars
        ]

        # Check for patterns with weighted scoring
        suspicious_count = 0
        for pattern in corruption_patterns:
            matches = re.findall(pattern, text)
            for match in matches:
                # Weight more severe patterns higher
                if pattern in [r'[\uFFFD]', r'[\x00-\x08\x0B\x0C\x0E-\x1F\x7F]']:
                    suspicious_count += len(match) * 2
                else:
                    suspicious_count += len(match)
        
        # Check for invalid Unicode sequences
        try:
            text.encode('utf-8').decode('utf-8')
        except UnicodeError:
            return True
            
        # Check for excessive repetition of the same character (e.g., 'kkkkkk')
        for i in range(0, len(text) - 5):
            if len(set(text[i:i+5])) == 1 and text[i] not in ' .-_,':
                suspicious_count += 10  # Increased weight for repetition
                break
        
        # Check for mixed encoding (e.g., UTF-8 interpreted as Windows-1252)
        try:
            text.encode('latin-1').decode('utf-8')
        except UnicodeError:
            pass
        else:
            # If it can be interpreted as both, it might be double-encoded
            suspicious_count += int(0.2 * len(text.strip()))  # Increased penalty
            
        # Check for excessive whitespace patterns
        if re.search(r'\s{10,}', text.strip()):
            suspicious_count += 20
            
        # Check for common corrupted text patterns
        if re.search(r'[\x00-\x08\x0B\x0C\x0E-\x1F\x7F]', text.strip()):
            suspicious_count += 15
        
        # Calculate ratio of suspicious content
        suspicious_ratio = suspicious_count / len(text.strip())
        self.logger.info(f"Suspicious character ratio: {suspicious_ratio:.3f} (threshold: {threshold})")
        
        return suspicious_ratio > threshold

    def _read_system_instructions(self, file_path: str = "prompts.yaml") -> dict:
        """
        Reads a YAML file and returns its content as a Python dictionary.
        """
        base_dir = os.path.dirname(os.path.abspath(__file__))
        instructions_path = os.path.join(base_dir, file_path)
        try:
            with open(instructions_path, "r", encoding="utf-8") as file:
                self.prompt_dictionary = yaml.safe_load(file)
        except FileNotFoundError:
            self.logger.error(f"Instructions file not found: '{instructions_path}'")
        except yaml.YAMLError as e:
            self.logger.error(f"Error parsing YAML file '{instructions_path}': {e}")
        except Exception as e:  # pylint: disable=broad-except
            self.logger.error(
                f"An unexpected error occurred while reading '{instructions_path}': {e}"
            )
        self.logger.info("Systems instructions loaded successfully.")

    def save_response_text(
        self,
        out_data: str,
        label: str = "response",
        idx: int = 1,
        chunk_size: int = 1024 * 1024,
        errors: str = "strict",
    ) -> str:
        """
        Saves a string variable to a text file with a dynamic filename.
        This method now incorporates the logic of the former _save_summary_txt function.
        Returns the full path to the saved file, or an empty string on error.
        """
        if out_data is None or out_data == "":
            return "No data to save"

        if label == "reference_summary":
            filename = f"{label}.{idx}.{self.run_tag}.txt"
        else:
            filename = f"{label}.{self.run_tag}.txt"
        filepath = os.path.join(self.output_directory, filename)

        try:
            with open(filepath, "w", encoding="utf-8", errors=errors) as f:
                for i in range(0, len(out_data), chunk_size):
                    f.write(out_data[i : i + chunk_size])
            return filepath
        except IOError as e:
            return f"Error saving file '{filepath}': {e}"

    def generate_tldr_pdf(self, summary_text, doc_title):
        """Saves polished summary string to formatted PDF document."""

        # Create file path with linted name
        file_path = self._create_filename(doc_title)

        # Format content
        styles = getSampleStyleSheet()
        h1_style = styles["h1"]
        h1_style.alignment = 1
        body = [
            Paragraph(doc_title.replace('"', ""), h1_style),
            Spacer(1, 0.15 * inch),
        ]
        body += self._interpret_markdown(summary_text)

        # Create document
        doc = SimpleDocTemplate(file_path, pagesize=letter)
        try:
            doc.build(body)
        except Exception as e:
            self.logger.error(f"An unexpected error occurred while saving {file_path}: {e}")
            raise

        return file_path

    @staticmethod
    def _interpret_markdown(text: str) -> list:
        """
        Converts custom markdown-like syntax to ReportLab-friendly HTML.
        - # Header 1      → <font size=18>
        - ## Header 2    → <font size=16>
        - ### Header 3  → <font size=14>
        - #### Header 4   → <font size=12>
        - *bold*          → <b>
        - ~italic~    → <i>
        Returns a list of Paragraph and Spacer elements.
        """
        styles = getSampleStyleSheet()
        normal = styles["Normal"]
        story = []

        # Header sizes mapping
        header_sizes = {1: 16, 2: 14, 3: 12, 4: 11}

        lines = text.strip().splitlines()

        for line in lines:
            line = line.strip()
            if not line:
                story.append(Spacer(1, 0.1 * inch))
                continue
            elif line.startswith("# "):
                continue
            elif line in ["###", "---"]:
                continue

            # Convert italic/bold (model uses *, shrugs) and subscripts
            line = re.sub(r"\*(.*?)\*", r"<b>\1</b>", line)
            line = re.sub(r"_(.*?)_", r"<sub>\1</sub>", line)

            # Header detection: one or more '#' followed by a space
            header_match = re.match(r"^(#{1,6})\s+(.*)", line)
            if header_match:
                level = len(header_match.group(1))
                content = header_match.group(2).strip()
                font_size = header_sizes.get(level, 11)  # default to 11 for > 4 #
                html_line = f'<b><font size="{font_size}">{content}</font></b>'
                story.append(Paragraph(html_line, normal))
                story.append(Spacer(1, 0.05 * inch))
            elif line.startswith("- "):
                bullet_content = line[2:].strip()
                html_line = f"• {bullet_content}"
                story.append(Paragraph(html_line, normal))
            else:
                html_line = line
                story.append(Paragraph(html_line, normal))

        return story

    @staticmethod
    def _create_filename(title: str, max_length: int = 50) -> str:
        """
        Accepts a string that is a document title, removes uninformative words,
        and reformats the string to be used as a file name.
        """
        # Convert to lowercase
        filename = title.lower()

        # Define uninformative words (can be extended)
        uninformative_words = [
            "a",
            "an",
            "the",
            "of",
            "in",
            "on",
            "at",
            "for",
            "with",
            "and",
            "or",
            "but",
            "is",
            "are",
            "was",
            "were",
            "be",
            "been",
            "being",
            "to",
            "from",
            "by",
            "as",
            "that",
            "which",
            "this",
            "these",
            "those",
            "it",
            "its",
            "about",
            "through",
            "beyond",
            "up",
            "down",
            "out",
            "into",
            "over",
            "under",
            "from",
            "around",
            "about",
            "via",
            "re",
            "regarding",
            "concerning",
            "document",
            "report",
            "summary",
            "efficient",
            "technical",
            "overview",
            "introduction",
            "advancements",
            "analysis",
            "study",
            "research",
            "paper",
            "article",
            "draft",
            "final",
            "version",
            "update",
            "notes",
            "memo",
            "brief",
            "presentation",
            "review",
            "whitepaper",
            "guide",
            "manual",
            "spec",
            "specification",
            "appendix",
            "chapter",
            "section",
            "part",
            "volume",
            "issue",
            "release",
            "plan",
            "project",
            "initiative",
            "program",
            "system",
            "process",
            "procedure",
            "framework",
            "methodology",
            "approach",
            "solution",
            "strategy",
            "tbd",
            "for review",
            "confidential",
        ]

        # Remove uninformative words
        for word in uninformative_words:
            filename = re.sub(r"\b" + re.escape(word) + r"\b", "", filename)

        # Replace non-alphanumeric characters (except spaces and underscores)
        filename = re.sub(r"[^a-z0-9\s_]", " ", filename)

        # Replace white space with underscore
        filename = "_".join(filename.split()).strip("_")

        # Truncate if too long
        if len(filename) > max_length:
            filename = filename[:max_length]
            # Try to avoid cutting off in the middle of a word if possible
            last_underscore = filename.rfind("_")
            if last_underscore != -1 and last_underscore > max_length - 20:
                filename = filename[:last_underscore]

        # Assemble final name
        return filename.strip("_") + ".tldr.pdf"
