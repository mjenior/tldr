import os
import re
import yaml
from datetime import datetime

from docx import Document
from PyPDF2 import PdfReader, PdfMerger
from bs4 import BeautifulSoup

from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.units import inch


def create_timestamp():
    """Generate a timestamp string (e.g., 20231027_103000)"""
    return datetime.now().strftime("%Y%m%d_%H%M%S")


def fetch_content(
    user_files: list = None,
    search_dir: str = ".",
    recursive: bool = False,
):
    """
    Find files and read in thier contents.
    Returns a dictionary with file extensions as keys by default.
    Also able to return all content in a single list.
    """

    # Collect readable text
    files = _find_readable_files(
        infiles=user_files, directory=search_dir, recursive=recursive
    )

    # Extract text from found files
    content = []
    for ext in files.keys():
        content += read_file_content(files[ext], ext)

    return content


def _find_readable_files(
    infiles: list = None, directory: str = ".", recursive=False
) -> dict:
    """
    Recursively scan the given directory for readable text files.
    Includes: .pdf, .docx, and .html files, and generic text files.
    Args:
        directory: The path to the directory to scan. Defaults to the current directory.
    Returns:
        A dictionary with file extensions as keys and lists of file paths as values.
    """

    readable_files_by_type = {"pdf": [], "docx": [], "html": [], "txt": [], "md": []}

    if infiles is not None:
        for file in infiles:
            ext = os.path.splitext(file)[1].lower()
            readable_files_by_type = _update_file_dictionary(
                readable_files_by_type, file, ext
            )
    else:
        if recursive == False:
            for filename in os.listdir(directory):
                if ".tldr." in filename:
                    continue
                else:
                    filepath = os.path.join(directory, filename)
                    ext = os.path.splitext(filename)[1].lower()

                readable_files_by_type = _update_file_dictionary(
                    readable_files_by_type, filepath, ext
                )

        else:
            for root, _, files in os.walk(directory):
                for filename in files:
                    if ".tldr." in filename:
                        continue
                    else:
                        filepath = os.path.join(root, filename)
                        ext = os.path.splitext(filename)[1].lower()

                    readable_files_by_type = _update_file_dictionary(
                        readable_files_by_type, filepath, ext
                    )

    return {k: v for k, v in readable_files_by_type.items() if v}


def _update_file_dictionary(file_dict, file_path, file_ext):
    """Add file content entries to content dictionary"""
    try:
        if file_ext in [".pdf"]:
            reader = PdfReader(file_path)
            # Check if any page seems to contain text
            if any(page.extract_text() for page in reader.pages):
                file_dict["pdf"].append(file_path)

        elif file_ext in [".docx"]:
            doc = Document(file_path)
            if any(para.text.strip() for para in doc.paragraphs):
                file_dict["docx"].append(file_path)

        elif file_ext in [".html", ".htm"]:
            with open(file_path, "r", encoding="utf-8") as f:
                f.read(1024)
            readable_files_by_type["html"].append(file_path)

        elif file_ext in [".md"]:
            with open(file_path, "r", encoding="utf-8") as f:
                f.read(1024)
            file_dict["md"].append(file_path)

        elif file_ext in [".txt"]:
            with open(file_path, "r", encoding="utf-8") as f:
                f.read(1024)
            file_dict["txt"].append(file_path)

    except Exception as e:
        pass

    return file_dict


def read_file_content(filelist, ext=None):
    """
    Reads the main body text from a given file paths and returns as strings.
    """
    if isinstance(filelist, str):
        filelist = [filelist]

    content_list = []
    for filepath in filelist:
        # If now extension is provided, try and grab it
        if ext == None:
            ext = filepath.split(".")[-1]

        try:
            if ext == "pdf":
                reader = PdfReader(filepath)
                text = ""
                for page in reader.pages:
                    text += page.extract_text() + "\n"
            elif "doc" in ext:
                doc = Document(filepath)
                text = ""
                for para in doc.paragraphs:
                    text += para.text + "\n"
            elif "htm" in ext:
                with open(filepath, "r", encoding="utf-8") as f:
                    html_content = f.read()
                soup = BeautifulSoup(html_content, "html.parser")
                text = soup.get_text(separator="\n")
            else:
                with open(filepath, "r", encoding="utf-8") as f:
                    text = f.read()
        except Exception as e:
            print(f"Error reading file '{filepath}': {e}")

        content_list.append(text.strip())

    return content_list


def read_system_instructions(file_path: str = "instructions.yaml") -> dict:
    """
    Reads a YAML file and returns its content as a Python dictionary.
    """
    tldr_path = os.path.dirname(os.path.abspath(__file__))
    instructions_path = os.path.join(tldr_path, file_path)
    try:
        with open(instructions_path, "r") as file:
            instructions = yaml.safe_load(file)
    except Exception as e:
        print(f"An unexpected error occurred while reading '{instructions_path}': {e}")

    return instructions


def save_response_text(
    out_data: str,
    label: str = "response",
    output_dir: str = ".",
    idx: int = 1,
) -> str:
    """
    Saves a large string variable to a text file with a dynamic filename
    """
    if label == "summary":
        filename = f"{label}.{idx}.tldr.{create_timestamp()}.txt"
    else:
        filename = f"{label}.tldr.{create_timestamp()}.txt"
    filepath = os.path.join(output_dir, filename)

    return _save_summary_txt(out_data, filepath)


def _save_summary_txt(text_data, outpath, errors="strict", chunk_size=1024 * 1024):
    """Save response text to .txt file"""
    try:
        with open(outpath, "w", encoding="utf-8", errors=errors) as f:
            # Write in chunks to avoid excessive memory usage for very large strings
            for i in range(0, len(text_data), chunk_size):
                f.write(text_data[i : i + chunk_size])
        return f"Intermediate text saved to {outpath}"
    except Exception as e:
        return f"An unexpected error occurred while saving {outpath}: {e}"


def generate_tldr_pdf(summary_text, doc_title):
    """Saves polished summary string to formatted PDF document."""

    # Create file path with linted name
    file_path = _create_filename(doc_title)

    # Format content
    styles = getSampleStyleSheet()
    h1_style = styles["h1"]
    h1_style.alignment = 1
    body = [
        Paragraph(doc_title.replace('"', ""), h1_style),
        Spacer(1, 0.15 * inch),
    ]
    body += _interpret_markdown(summary_text)

    # Create document
    doc = SimpleDocTemplate(file_path, pagesize=letter)
    try:
        doc.build(body)
    except Exception as e:
        print(f"An unexpected error occurred while saving {file_path}: {e}")
        raise

    return file_path


def _interpret_markdown(text: str) -> list:
    """
    Converts custom markdown-like syntax to ReportLab-friendly HTML.
    - # Header 1      → <font size=18>
    - ## Header 2    → <font size=16>
    - ### Header 3  → <font size=14>
    - #### Header 4   → <font size=12>
    - *bold*          → <b>
    - ~italic~    → <i>
    Returns a list of Paragraph and Spacer elements.
    """
    styles = getSampleStyleSheet()
    normal = styles["Normal"]
    story = []

    # Header sizes mapping
    header_sizes = {1: 16, 2: 14, 3: 12, 4: 11}

    lines = text.strip().splitlines()

    for line in lines:
        line = line.strip()
        if not line:
            story.append(Spacer(1, 0.1 * inch))
            continue
        elif line.startswith("# "):
            continue
        elif line in ["###", "---"]:
            continue

        # Convert italic/bold (model uses *, shrugs) and subscripts
        line = re.sub(r"\*(.*?)\*", r"<b>\1</b>", line)
        line = re.sub(r"_(.*?)_", r"<sub>\1</sub>", line)

        # Header detection: one or more '#' followed by a space
        header_match = re.match(r"^(#{1,6})\s+(.*)", line)
        if header_match:
            level = len(header_match.group(1))
            content = header_match.group(2).strip()
            font_size = header_sizes.get(level, 11)  # default to 11 for > 4 #
            html_line = f'<b><font size="{font_size}">{content}</font></b>'
            story.append(Paragraph(html_line, normal))
            story.append(Spacer(1, 0.05 * inch))
        elif line.startswith("- "):
            bullet_content = line[2:].strip()
            html_line = f"• {bullet_content}"
            story.append(Paragraph(html_line, normal))
        else:
            html_line = line
            story.append(Paragraph(html_line, normal))

    return story


def _create_filename(title: str, max_length: int = 50) -> str:
    """
    Accepts a string that is a document title, removes uninformative words,
    and reformats the string to be used as a file name.
    """
    # Convert to lowercase
    filename = title.lower()

    # Define uninformative words (can be extended)
    uninformative_words = [
        "a",
        "an",
        "the",
        "of",
        "in",
        "on",
        "at",
        "for",
        "with",
        "and",
        "or",
        "but",
        "is",
        "are",
        "was",
        "were",
        "be",
        "been",
        "being",
        "to",
        "from",
        "by",
        "as",
        "that",
        "which",
        "this",
        "these",
        "those",
        "it",
        "its",
        "about",
        "through",
        "beyond",
        "up",
        "down",
        "out",
        "into",
        "over",
        "under",
        "from",
        "around",
        "about",
        "via",
        "re",
        "regarding",
        "concerning",
        "document",
        "report",
        "summary",
        "efficient",
        "technical",
        "overview",
        "introduction",
        "advancements",
        "analysis",
        "study",
        "research",
        "paper",
        "article",
        "draft",
        "final",
        "version",
        "update",
        "notes",
        "memo",
        "brief",
        "presentation",
        "review",
        "whitepaper",
        "guide",
        "manual",
        "spec",
        "specification",
        "appendix",
        "chapter",
        "section",
        "part",
        "volume",
        "issue",
        "release",
        "plan",
        "project",
        "initiative",
        "program",
        "system",
        "process",
        "procedure",
        "framework",
        "methodology",
        "approach",
        "solution",
        "strategy",
        "tbd",
        "for review",
        "confidential",
    ]

    # Remove uninformative words
    for word in uninformative_words:
        filename = re.sub(r"\b" + re.escape(word) + r"\b", "", filename)

    # Replace non-alphanumeric characters (except spaces and underscores)
    filename = re.sub(r"[^a-z0-9\s_]", " ", filename)

    # Replace white space with underscore
    filename = "_".join(filename.split()).strip("_")

    # Truncate if too long
    if len(filename) > max_length:
        filename = filename[:max_length]
        # Try to avoid cutting off in the middle of a word if possible
        last_underscore = filename.rfind("_")
        if last_underscore != -1 and last_underscore > max_length - 20:
            filename = filename[:last_underscore]

    # Assemble final name
    return filename.strip("_") + ".tldr.pdf"
