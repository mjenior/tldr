
import os
import yaml
from datetime import datetime

from docx import Document
from PyPDF2 import PdfReader, PdfMerger
from bs4 import BeautifulSoup


def fetch_content(search_dir):
	"""
	Find files and read in thier contents. 
	Returns a dictionary with file extensions as keys.
	"""

	# Collect readable text
	files = _find_readable_files(search_dir)

	# Extract text from found files
	content = {}
	for ext in files.keys():
		content[ext] = _read_file_content(files[ext], ext)

	return content


def _find_readable_files(directory: str = '.') -> dict:
    """
    Recursively scan the given directory for readable text files.
    Includes: .pdf, .docx, and .html files, and generic text files.
    Args:
        directory: The path to the directory to scan. Defaults to the current directory.
    Returns:
        A dictionary with file extensions as keys and lists of file paths as values.
    """
    readable_files_by_type = {'pdf': [], 'docx': [], 'html': [], 'txt': []}

    for root, _, files in os.walk(directory):
        for filename in files:
            if filename.startswith('tldr.') == True:
                continue

            filepath = os.path.join(root, filename)
            ext = os.path.splitext(filename)[1].lower()

            try:
                if ext in ['.pdf']:
                    reader = PdfReader(filepath)
                    # Check if any page seems to contain text
                    if any(page.extract_text() for page in reader.pages):
                        readable_files_by_type['pdf'].append(filepath)

                elif ext in ['.docx']:
                    doc = Document(filepath)
                    if any(para.text.strip() for para in doc.paragraphs):
                        readable_files_by_type['docx'].append(filepath)

                elif ext in ['.html', '.htm']:
                    with open(filepath, 'r', encoding='utf-8') as f:
                        f.read(1024)
                    readable_files_by_type['html'].append(filepath)

                elif ext in ['.txt']:
                    with open(filepath, 'r', encoding='utf-8') as f:
                        f.read(1024)
                    readable_files_by_type['txt'].append(filepath)

            except Exception as e:
                continue

    return {k: v for k, v in readable_files_by_type.items() if v}


def _read_file_content(filelist, ext):
    """
    Reads the main body text from a given file paths and returns as strings.
    """
    content_list = []
    for filepath in filelist:

        try:
            if ext == 'pdf':
                reader = PdfReader(filepath)
                text = ""
                for page in reader.pages:
                    text += page.extract_text() + "\n"
            elif 'doc' in ext:
                doc = Document(filepath)
                text = ""
                for para in doc.paragraphs:
                    text += para.text + "\n"
            elif 'htm' in ext:
                with open(filepath, 'r', encoding='utf-8') as f:
                    html_content = f.read()
                soup = BeautifulSoup(html_content, 'html.parser')
                text = soup.get_text(separator='\n')
            else:
                with open(filepath, 'r', encoding='utf-8') as f:
                    text = f.read()
        except Exception as e:
            print(f"Error reading file '{filepath}': {e}")

        content_list.append(text.strip())

    return content_list


def read_system_instructions(file_path: str = "instructions.yaml") -> dict:
    """
    Reads a YAML file and returns its content as a Python dictionary.
    """
    tldr_path = os.path.dirname(os.path.abspath(__file__))
    instructions_path = os.path.join(tldr_path, file_path)
    try:
        with open(instructions_path, "r") as file:
            data = yaml.safe_load(file)
            if isinstance(data, dict):
                instructions = data
            else:
                print(
                    f"Error: Content in '{instructions_path}' is not a dictionary (YAML root is type: {type(data)})."
                )
    except Exception as e:
        print(f"An unexpected error occurred while reading '{instructions_path}': {e}")

    return instructions


def save_response_text(
    data_str: str, label: str = "response", output_dir: str = "."
) -> str:
    """
    Saves a large string variable to a text file with a dynamic filename
    based on a timestamp and a user-provided label.
    """
    errors = "strict"
    chunk_size = 1024 * 1024

    # Generate a timestamp string (e.g., 20231027_103000)
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")

    # Sanitize the label to create a valid filename part
    # Replace spaces with underscores and remove characters not suitable for filenames
    sanitized_label = "".join(
        c if c.isalnum() or c in ("_", "-") else "_" for c in label
    ).strip("_")

    # Construct the dynamic filename
    filename = f"tldr.{sanitized_label}.{timestamp}.txt"
    filepath = os.path.join(output_dir, filename)

    # Format text just in case
    if isinstance(data_str, list):
        data_str = "\n\n".join([str(x) for x in data_str])
    else:
        data_str = str(data_str)

    try:
        with open(filename, "w", encoding="utf-8", errors=errors) as f:
            # Write in chunks to avoid excessive memory usage for very large strings
            for i in range(0, len(data_str), chunk_size):
                # Note: Corrected variable name from data_string to data_str
                f.write(data_str[i : i + chunk_size])

        print(f"\tSaved data to {filename}")

    except IOError as e:
        print(f"Error writing to file {filename}: {e}")
        raise  # Re-raise the exception after printing
    except Exception as e:
        print(f"An unexpected error occurred while saving {filename}: {e}")
        raise  # Re-raise the exception
